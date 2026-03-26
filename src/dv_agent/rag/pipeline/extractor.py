"""
Document Content Extractor
文档内容提取器

基于 Unstructured 库提取多格式文档内容。
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional, Union

from .detector import DocumentDetector, FileFormat

logger = logging.getLogger(__name__)


@dataclass
class ExtractedElement:
    """提取的文档元素"""
    
    type: str  # text, title, table, image, etc.
    content: str
    metadata: dict = field(default_factory=dict)
    page_number: Optional[int] = None


@dataclass
class ExtractedContent:
    """提取的文档内容"""
    
    elements: list[ExtractedElement] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    errors: list[str] = field(default_factory=list)
    
    @property
    def text(self) -> str:
        """获取纯文本内容（合并所有元素）"""
        return "\n\n".join(
            elem.content for elem in self.elements
            if elem.content.strip()
        )
    
    @property
    def page_count(self) -> Optional[int]:
        """获取页数"""
        if not self.elements:
            return None
        
        page_numbers = [
            elem.page_number for elem in self.elements
            if elem.page_number is not None
        ]
        
        return max(page_numbers) if page_numbers else None
    
    @property
    def has_errors(self) -> bool:
        """是否有错误"""
        return len(self.errors) > 0


class DocumentExtractor:
    """
    文档内容提取器
    
    使用 Unstructured 库提取多格式文档内容。
    支持 PDF、Word、Excel、PPT、HTML、Markdown 等格式。
    """
    
    def __init__(
        self,
        ocr_languages: Optional[list[str]] = None,
        strategy: str = "auto",
        use_hi_res: bool = False,
    ):
        """
        初始化提取器
        
        Args:
            ocr_languages: OCR 识别语言列表，如 ["chi_sim", "eng"]
            strategy: 提取策略 ("auto", "fast", "hi_res", "ocr_only")
            use_hi_res: 是否使用高精度模式（较慢但更准确）
        """
        self.ocr_languages = ocr_languages or ["chi_sim", "eng"]
        self.strategy = "hi_res" if use_hi_res else strategy
        self.detector = DocumentDetector()
        
        self._unstructured_available = self._check_unstructured()
    
    def _check_unstructured(self) -> bool:
        """检查 Unstructured 库是否可用"""
        try:
            from unstructured.partition.auto import partition
            return True
        except ImportError:
            logger.warning(
                "unstructured not installed. "
                "Run: pip install 'unstructured[all-docs]'"
            )
            return False
    
    def extract(
        self,
        file_path: Optional[Union[str, Path]] = None,
        content: Optional[bytes] = None,
        filename: Optional[str] = None,
        file_format: Optional[FileFormat] = None,
    ) -> ExtractedContent:
        """
        提取文档内容
        
        Args:
            file_path: 文件路径
            content: 文件内容（字节）
            filename: 文件名
            file_format: 文件格式（可选，自动检测）
            
        Returns:
            提取的内容
        """
        result = ExtractedContent()
        
        # 检测文件格式
        if file_format is None:
            file_format = self.detector.detect(file_path, content, filename)
        
        if not file_format.is_supported:
            result.errors.append(
                f"Unsupported file format: {file_format.value}. "
                f"Supported formats: {self.detector.get_supported_formats()}"
            )
            return result
        
        # 选择提取方法
        try:
            if self._unstructured_available:
                return self._extract_with_unstructured(
                    file_path, content, filename, file_format
                )
            else:
                return self._extract_fallback(
                    file_path, content, filename, file_format
                )
        except Exception as e:
            logger.error(f"Extraction failed: {e}")
            result.errors.append(f"Extraction failed: {str(e)}")
            return result
    
    def _extract_with_unstructured(
        self,
        file_path: Optional[Union[str, Path]],
        content: Optional[bytes],
        filename: Optional[str],
        file_format: FileFormat,
    ) -> ExtractedContent:
        """
        使用 Unstructured 提取内容
        """
        from unstructured.partition.auto import partition
        
        result = ExtractedContent()
        
        # 准备参数
        kwargs: dict[str, Any] = {
            "strategy": self.strategy,
        }
        
        # 添加 OCR 语言配置
        if file_format == FileFormat.PDF:
            kwargs["languages"] = self.ocr_languages
        
        # 执行提取
        try:
            if file_path:
                elements = partition(filename=str(file_path), **kwargs)
            elif content and filename:
                from io import BytesIO
                elements = partition(
                    file=BytesIO(content),
                    metadata_filename=filename,
                    **kwargs
                )
            else:
                result.errors.append("No file path or content provided")
                return result
        except Exception as e:
            result.errors.append(f"Partition failed: {str(e)}")
            return result
        
        # 转换元素
        for elem in elements:
            extracted_elem = ExtractedElement(
                type=elem.category if hasattr(elem, "category") else "text",
                content=str(elem),
                metadata=elem.metadata.to_dict() if hasattr(elem, "metadata") else {},
                page_number=getattr(elem.metadata, "page_number", None) if hasattr(elem, "metadata") else None,
            )
            result.elements.append(extracted_elem)
        
        # 提取元数据
        if elements and hasattr(elements[0], "metadata"):
            first_meta = elements[0].metadata
            result.metadata = {
                "filename": getattr(first_meta, "filename", filename),
                "filetype": getattr(first_meta, "filetype", file_format.value),
            }
        
        return result
    
    def _extract_fallback(
        self,
        file_path: Optional[Union[str, Path]],
        content: Optional[bytes],
        filename: Optional[str],
        file_format: FileFormat,
    ) -> ExtractedContent:
        """
        备用提取方法（不依赖 Unstructured）
        """
        result = ExtractedContent()
        
        # 获取内容
        if content is None and file_path:
            with open(file_path, "rb") as f:
                content = f.read()
        
        if content is None:
            result.errors.append("No content to extract")
            return result
        
        # 根据格式选择提取方法
        if file_format == FileFormat.PDF:
            return self._extract_pdf_fallback(content, result)
        elif file_format in {FileFormat.TXT, FileFormat.MD}:
            return self._extract_text_fallback(content, result)
        elif file_format in {FileFormat.HTML, FileFormat.HTM}:
            return self._extract_html_fallback(content, result)
        elif file_format == FileFormat.DOCX:
            return self._extract_docx_fallback(content, result)
        else:
            result.errors.append(
                f"Fallback extraction not available for {file_format.value}. "
                "Please install unstructured."
            )
            return result
    
    def _extract_pdf_fallback(
        self,
        content: bytes,
        result: ExtractedContent,
    ) -> ExtractedContent:
        """使用 PyMuPDF 提取 PDF 内容"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=content, filetype="pdf")
            
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                if text.strip():
                    result.elements.append(ExtractedElement(
                        type="text",
                        content=text,
                        page_number=page_num,
                    ))
            
            result.metadata["page_count"] = len(doc)
            doc.close()
            
        except ImportError:
            result.errors.append(
                "PyMuPDF not installed. Run: pip install PyMuPDF"
            )
        except Exception as e:
            result.errors.append(f"PDF extraction failed: {str(e)}")
        
        return result
    
    def _extract_text_fallback(
        self,
        content: bytes,
        result: ExtractedContent,
    ) -> ExtractedContent:
        """提取纯文本内容"""
        try:
            # 尝试多种编码
            for encoding in ["utf-8", "gbk", "gb2312", "latin1"]:
                try:
                    text = content.decode(encoding)
                    result.elements.append(ExtractedElement(
                        type="text",
                        content=text,
                    ))
                    break
                except UnicodeDecodeError:
                    continue
            else:
                result.errors.append("Failed to decode text content")
        except Exception as e:
            result.errors.append(f"Text extraction failed: {str(e)}")
        
        return result
    
    def _extract_html_fallback(
        self,
        content: bytes,
        result: ExtractedContent,
    ) -> ExtractedContent:
        """使用 BeautifulSoup 提取 HTML 内容"""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(content, "html.parser")
            
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator="\n")
            
            # 清理空白行
            lines = [line.strip() for line in text.splitlines()]
            text = "\n".join(line for line in lines if line)
            
            if text:
                result.elements.append(ExtractedElement(
                    type="text",
                    content=text,
                ))
            
            # 提取标题
            title = soup.find("title")
            if title:
                result.metadata["title"] = title.get_text()
            
        except ImportError:
            result.errors.append(
                "BeautifulSoup not installed. Run: pip install beautifulsoup4"
            )
        except Exception as e:
            result.errors.append(f"HTML extraction failed: {str(e)}")
        
        return result
    
    def _extract_docx_fallback(
        self,
        content: bytes,
        result: ExtractedContent,
    ) -> ExtractedContent:
        """使用 python-docx 提取 Word 内容"""
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(content))
            
            for para in doc.paragraphs:
                if para.text.strip():
                    result.elements.append(ExtractedElement(
                        type="paragraph",
                        content=para.text,
                    ))
            
            # 提取表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                
                if table_text:
                    result.elements.append(ExtractedElement(
                        type="table",
                        content="\n".join(table_text),
                    ))
            
        except ImportError:
            result.errors.append(
                "python-docx not installed. Run: pip install python-docx"
            )
        except Exception as e:
            result.errors.append(f"DOCX extraction failed: {str(e)}")
        
        return result
